from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status
from sqlalchemy.orm import Session
from typing import List, Optional
from datetime import datetime
import os
import shutil

from ....services.rag import index_document, retrieve_relevant_chunks, delete_document_chunks
from ....core.database import get_db                 # ✅ 4 dots
from ....core.security import get_current_user       # ✅ 4 dots + security (not auth)
from ....models.user import User                     # ✅ 4 dots
from ....models.document import Document, DocumentStatus
from ....models.chat import Chat, ChatSession
from ....schemas.document import DocumentResponse, ChatRequest, ChatResponse, ChatSessionResponse, ChatSessionDetail, TranslateRequest, TranslateResponse
from ....services.gemini import generate_summary, chat_with_document, translate_summary, analyze_document
from ....core.config import settings

router = APIRouter(prefix="/documents", tags=["documents"])

ALLOWED_EXTENSIONS = ['pdf', 'docx', 'jpg', 'jpeg', 'png']

@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    model_choice: str = Form("flash"),
    language: str = Form("English"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    ext = file.filename.split('.')[-1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"Unsupported file type: {ext}")
    
    upload_dir = "uploads"
    os.makedirs(upload_dir, exist_ok=True)
    file_path = os.path.join(upload_dir, f"{datetime.now().timestamp()}_{file.filename}")
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    
    # Extract text
    try:
        if ext == 'pdf':
            text = extract_text_from_pdf(file_path)
        elif ext == 'docx':
            text = extract_text_from_docx(file_path)
        elif ext in ['jpg', 'jpeg', 'png']:
            text = extract_text_from_image(file_path)
        else:
            text = ""
    except Exception as e:
        text = f"[Error extracting text: {str(e)}]"
    
    summary = generate_summary(text, model_choice, language)
    analysis = analyze_document(text, model_choice)  # ✅ naya smart analysis

    db_doc = Document(
        filename=file.filename,
        file_path=file_path,
        extracted_text=text,
        summary=summary,
        model_used=model_choice,
        ocr_lang=language,
        status=DocumentStatus.COMPLETED,
        user_id=current_user.id,
        created_at=datetime.utcnow(),
        risk_score=analysis.get("risk_score"),
        risk_level=analysis.get("risk_level"),
        important_clauses=analysis.get("important_clauses"),
        missing_clauses=analysis.get("missing_clauses"),
        pii_findings=analysis.get("pii_findings"),
        pii_found_count=analysis.get("pii_found_count"),
        recommendations=analysis.get("recommendations"),
        suggested_questions=analysis.get("suggested_questions"),
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)

    index_document(db_doc.id, current_user.id, text)   # ✅ RAG: chunk + embed + store
    
    return DocumentResponse(
        id=db_doc.id,
        filename=db_doc.filename,
        summary=db_doc.summary,
        status=db_doc.status,
        model_used=db_doc.model_used,
        language=db_doc.ocr_lang,
        created_at=db_doc.created_at,
        risk_score=db_doc.risk_score,
        risk_level=db_doc.risk_level,
        important_clauses=db_doc.important_clauses,
        missing_clauses=db_doc.missing_clauses,
        pii_findings=db_doc.pii_findings,
        pii_found_count=db_doc.pii_found_count,
        recommendations=db_doc.recommendations,
        suggested_questions=db_doc.suggested_questions,
    )

@router.get("/", response_model=List[DocumentResponse])
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    docs = db.query(Document).filter(Document.user_id == current_user.id).order_by(Document.created_at.desc()).all()
    return docs

@router.get("/{doc_id}", response_model=DocumentResponse)
async def get_document(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(404, "Document not found")

    return DocumentResponse(
        id=doc.id,
        filename=doc.filename,
        summary=doc.summary,
        status=doc.status,
        model_used=doc.model_used,
        language=doc.ocr_lang,
        created_at=doc.created_at,
        risk_score=doc.risk_score,
        risk_level=doc.risk_level,
        important_clauses=doc.important_clauses,
        missing_clauses=doc.missing_clauses,
        pii_findings=doc.pii_findings,
        pii_found_count=doc.pii_found_count,
        recommendations=doc.recommendations,
        suggested_questions=doc.suggested_questions,
    )

@router.delete("/{doc_id}")
async def delete_document(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(404, "Document not found")
    try:
        if doc.file_path and os.path.exists(doc.file_path):
            os.remove(doc.file_path)
    except:
        pass
    delete_document_chunks(doc.id)   # ✅ RAG cleanup
    db.delete(doc)
    db.commit()
    return {"message": "Document deleted successfully"}

@router.post("/{doc_id}/chat", response_model=ChatResponse)
async def chat(
    doc_id: int,
    request: ChatRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(404, "Document not found")

    if request.session_id:
        session = db.query(ChatSession).filter(
            ChatSession.id == request.session_id,
            ChatSession.document_id == doc_id,
            ChatSession.user_id == current_user.id,
        ).first()
        if not session:
            raise HTTPException(404, "Chat session not found")
    else:
        title = request.question.strip()
        if len(title) > 60:
            title = title[:57] + "..."
        session = ChatSession(document_id=doc_id, user_id=current_user.id, title=title)
        db.add(session)
        db.commit()
        db.refresh(session)

    # ✅ RAG: retrieve only the most relevant chunks instead of sending the whole document
    relevant_chunks = retrieve_relevant_chunks(doc.id, request.question, top_k=5)
    context = "\n\n---\n\n".join(relevant_chunks) if relevant_chunks else (doc.extracted_text or "")

    answer = chat_with_document(request.question, context, request.model_choice)

    chat_record = Chat(
        session_id=session.id,
        user_id=current_user.id,
        document_id=doc.id,
        question=request.question,
        answer=answer,
    )
    db.add(chat_record)
    db.commit()
    db.refresh(chat_record)

    return ChatResponse(
        question=chat_record.question,
        answer=chat_record.answer,
        created_at=chat_record.created_at,
        session_id=session.id,
        session_title=session.title,
    )

@router.get("/{doc_id}/chat/sessions", response_model=List[ChatSessionResponse])
async def get_chat_sessions(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(404, "Document not found")
    sessions = (
        db.query(ChatSession)
        .filter(ChatSession.document_id == doc_id, ChatSession.user_id == current_user.id)
        .order_by(ChatSession.created_at.desc())
        .all()
    )
    return sessions

@router.get("/{doc_id}/chat/sessions/{session_id}", response_model=ChatSessionDetail)
async def get_chat_session_detail(
    doc_id: int,
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    session = db.query(ChatSession).filter(
        ChatSession.id == session_id,
        ChatSession.document_id == doc_id,
        ChatSession.user_id == current_user.id,
    ).first()
    if not session:
        raise HTTPException(404, "Chat session not found")

    messages = (
        db.query(Chat)
        .filter(Chat.session_id == session_id, Chat.user_id == current_user.id)
        .order_by(Chat.created_at.asc())
        .all()
    )

    return ChatSessionDetail(
        id=session.id,
        title=session.title,
        created_at=session.created_at,
        messages=messages,
    )

@router.delete("/{doc_id}/chat/sessions/{session_id}")
async def delete_chat_session(
    doc_id: int,
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    session = db.query(ChatSession).filter(
        ChatSession.id == session_id,
        ChatSession.document_id == doc_id,
        ChatSession.user_id == current_user.id,
    ).first()
    if not session:
        raise HTTPException(404, "Chat session not found")

    db.query(Chat).filter(Chat.session_id == session_id).delete()
    db.delete(session)
    db.commit()
    return {"message": "Chat session deleted"}

@router.delete("/{doc_id}/chat/sessions")
async def clear_all_chat_sessions(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(404, "Document not found")
    session_ids = [
        s.id for s in db.query(ChatSession)
        .filter(ChatSession.document_id == doc_id, ChatSession.user_id == current_user.id)
        .all()
    ]
    db.query(Chat).filter(Chat.session_id.in_(session_ids)).delete(synchronize_session=False)
    db.query(ChatSession).filter(ChatSession.id.in_(session_ids)).delete(synchronize_session=False)
    db.commit()
    return {"message": "All chat sessions cleared"}

@router.post("/{doc_id}/translate", response_model=TranslateResponse)
async def translate_document_summary(
    doc_id: int,
    request: TranslateRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(404, "Document not found")
    translated = translate_summary(doc.summary, request.language, "flash")
    return TranslateResponse(
        summary=translated,
        language=request.language,
    )

@router.get("/dashboard/stats")
async def get_dashboard_stats(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    total = db.query(Document).filter(Document.user_id == current_user.id).count()
    analyzed = db.query(Document).filter(
        Document.user_id == current_user.id,
        Document.status == DocumentStatus.COMPLETED
    ).count()
    recent_docs = db.query(Document).filter(
        Document.user_id == current_user.id,
        Document.status == DocumentStatus.COMPLETED
    ).order_by(Document.created_at.desc()).limit(5).all()
    return {
        "total_documents": total,
        "documents_analyzed": analyzed,
        "average_risk_score": None,
        "ai_queries": 0,
        "recent_documents": recent_docs,
    }

# Helper functions
def extract_text_from_pdf(file_path: str) -> str:
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ''.join([page.extract_text() or '' for page in reader.pages])

        # Agar embedded text bahut kam/khaali mila (scanned/image-based PDF), OCR fallback try karo
        if len(text.strip()) < 20:
            text = extract_text_from_scanned_pdf(file_path)

        return text
    except ImportError:
        return "[PyPDF2 not installed]"
    except Exception as e:
        return f"[Error: {str(e)}]"

def extract_text_from_scanned_pdf(file_path: str) -> str:
    try:
        from pdf2image import convert_from_path
        import pytesseract

        if settings.TESSERACT_PATH:
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_PATH

        pages = convert_from_path(file_path, poppler_path=settings.POPPLER_PATH or None)
        text = "\n".join(
            pytesseract.image_to_string(page, lang="eng+hin+mar") for page in pages
        )
        return text
    except ImportError:
        return "[pdf2image or pytesseract not installed — cannot OCR scanned PDF]"
    except Exception as e:
        return f"[Error OCR-ing scanned PDF: {str(e)}]"

def extract_text_from_docx(file_path: str) -> str:
    try:
        import docx
        doc = docx.Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    except ImportError:
        return "[python-docx not installed]"
    except Exception as e:
        return f"[Error: {str(e)}]"

def extract_text_from_image(file_path: str) -> str:
    try:
        from PIL import Image
        import pytesseract
        if settings.TESSERACT_PATH:
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_PATH
        # eng+hin+mar: Tesseract in teeno languages ko ek saath try karega,
        # isliye chahe document English/Hindi/Marathi kisi me bhi ho (ya mix ho), extract ho jayega.
        return pytesseract.image_to_string(Image.open(file_path), lang="eng+hin+mar")
    except ImportError:
        return "[Pillow or pytesseract not installed]"
    except Exception as e:
        return f"[Error: {str(e)}]"